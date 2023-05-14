from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
letra=None
tamaño=None
def agregar_portada(titulo,estudiante,universidad,facultad,materia,docente,fecha,letra,tamaño):
    document = Document()
    # Tipo de letra y tamaño
    style = document.styles['Normal']
    font = style.font
    font.name = letra
    font.size = Pt(tamaño)

    # Interlineado
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 2
    # Margen
    sections = document.sections
    for section in sections:
        i=0
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = f"{titulo}\t\t {i+1}"
        paragraph.style = document.styles["Header"]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    document.add_paragraph(" ")
    title = document.add_paragraph()
    title_run = title.add_run(titulo)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(" ")
    for estudiante_singular in estudiante:
        document.add_paragraph(estudiante_singular[0]+".")
        document.add_paragraph(estudiante_singular[1]+".")
    document.add_paragraph(" ")
    document.add_paragraph(universidad+".")
    document.add_paragraph("Facultad de "+facultad+".")
    document.add_paragraph(materia+".")
    document.add_paragraph(docente+".")
    document.add_paragraph(fecha+".")
    for paragraph in document.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(f'docx/{titulo}.docx')
    print(f"""

    Se ha creado el documento "{titulo}.docx" de forma exitosa.
    
    """)
def capitalizar_mayores_tres(txt):
    resultado = txt.split()
    for i in range(len(resultado)):
        if len(resultado[i]) > 3:
            resultado[i] = resultado[i].capitalize()
    return " ".join(resultado)
def ejecucion():
    fuentes=[["Calibri",11],["Arial",11],["Lucida Sans Unicode",10],["Times New Roman",12],["Georgia",11],["Computer Modern",10]]
    print("""
    1- Calibri, 11
    2- Arial, 11
    3- Lucida Sans Unicode, 10
    4- Times New Roman, 12
    5- Georgia, 11
    6- Computer Modern, 10
    """)
    fuente= int(input("Tipo de fuente: "))-1
    fuente = fuentes[fuente]
    letra=fuente[0]
    tamaño=fuente[1]
    titulo=capitalizar_mayores_tres(str(input("Titulo: ")))
    n_est = int(input("Numero de integrantes: "))
    estudiante=[]
    for i in range (1,n_est+1):
        nombre_est=capitalizar_mayores_tres(str(input("Nombre del estudiante: ")))
        codigo_est=capitalizar_mayores_tres(str(input("Codigo del estudiante: ")))
        estudiante.append([nombre_est,codigo_est])
    estudiante=sorted(estudiante, key=lambda x: x[0])
    universidad=capitalizar_mayores_tres("Universidad tecnológica de bolívar")
    
    facultades=["Ingenieria","Ciencias Basicas","Ciencias Sociales y Humanidades","Arquitectura","Derecho"]
    print("""Facultades:
1- Ingeniería
2- Ciencias Básicas
3- Ciencias Sociales y Humanidades
4- Arquitectura
5- Derecho""")
    facultad=int(input("Indique la facultad: "))-1
    facultad=facultades[facultad]
    materia=capitalizar_mayores_tres(str(input("Materia: ")))
    docente=capitalizar_mayores_tres(str(input("Docente: ")))
    fecha=input("Dia de entrega: ")+" de "+input("Mes de entrega: ")+" de "+input("Año: ")
    agregar_portada(titulo,estudiante,universidad,facultad,materia,docente,fecha,letra,tamaño)
try:
    ejecucion()
except Exception as e:
    print(f"Ha ocurrido un error, codigo: {e}")
